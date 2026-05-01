"""Extract plain text from common document formats (PDF, DOCX, plain)."""

from __future__ import annotations

import io
from typing import Tuple

from .logging import get_logger

logger = get_logger(__name__)


def extract_text(raw: bytes, media_type: str, filename: str) -> Tuple[str, bool]:
    """Return (text, ok). ok=False means we couldn't extract anything useful.

    Supports application/pdf, docx (officedocument), plaintext-ish types. For
    everything else falls back to a UTF-8 decode-or-empty.
    """
    mt = (media_type or "").lower()
    name = (filename or "").lower()

    try:
        if mt == "application/pdf" or name.endswith(".pdf"):
            return _extract_pdf(raw), True

        if (
            mt in ("application/vnd.openxmlformats-officedocument.wordprocessingml.document",)
            or name.endswith(".docx")
        ):
            return _extract_docx(raw), True

        # Plaintext-ish — text/*, JSON, CSV, YAML, markdown, etc.
        if mt.startswith("text/") or mt in (
            "application/json",
            "application/x-yaml",
            "application/yaml",
        ) or name.endswith((".md", ".txt", ".csv", ".json", ".yaml", ".yml", ".log")):
            return raw.decode("utf-8", errors="replace"), True

    except Exception as e:
        logger.warning("Text extraction failed", filename=filename, media_type=mt, error=str(e))

    # Best-effort fallback for everything else
    try:
        return raw.decode("utf-8", errors="replace"), False
    except Exception:
        return "", False


def _extract_pdf(raw: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(raw))
    parts: list[str] = []
    for i, page in enumerate(reader.pages):
        try:
            t = page.extract_text() or ""
        except Exception:
            t = ""
        if t.strip():
            parts.append(f"--- Page {i + 1} ---\n{t.strip()}")
    return _sanitize(("\n\n".join(parts)))


def _extract_docx(raw: bytes) -> str:
    from docx import Document  # type: ignore

    doc = Document(io.BytesIO(raw))
    paragraphs = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
    # Tables
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))
    return _sanitize("\n".join(paragraphs))


def _sanitize(text: str) -> str:
    """Strip NUL bytes and other characters Postgres TEXT columns reject."""
    if not text:
        return ""
    # Postgres rejects \x00; replace common control chars except \t \n \r
    return "".join(
        ch for ch in text
        if ch in ("\t", "\n", "\r") or ord(ch) >= 0x20
    )
